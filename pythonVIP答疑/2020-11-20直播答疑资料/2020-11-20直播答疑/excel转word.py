# 教育机构 ：马士兵教育
# 讲    师：杨淑娟
#操作Excel的模块有 xlwt  ,xlwriter  ,xlrd,pandas ...
import openpyxl   #这是一个第三方模块，需要安装才能使用   操作Excel文件读写使用的
import docx
from docx.shared import RGBColor
#(1)将Excel文件中的内容读取到Python程序中来
wk=openpyxl.load_workbook('data/附件2：反洗钱人员考试题库.xlsx')  #相当于双击打开Excel文件
sheet=wk['通用题库']   #根据名称获取指定的sheet页  （工作表）
      #wk.active   获取当前的活动表
rows=sheet.rows     #获取所有数据行
lst=[]
for row in rows:   #遍历所有的行
    #每一行中有N多个单元格组成
    sub_lst=[]  #列表，用于存储一行中的N多个单元格的值
    for cell in row :   # 分别获取每行的单元格
        sub_lst.append(cell.value)   #添加的单元格的值
    lst.append(sub_lst)    #添加的是一行的数据
#
for item in lst:
    print(item)
# 转成Word文档

doc=docx.Document()  #创建一个文档对象
# 添加标题
doc.add_heading('反洗钱人员考试题库',level=1) #一级标题
#添加段落
para=doc.add_paragraph()  #添加一个空行  （标题和题干之间的空行）
for index in range(2,len(lst)):
    para=doc.add_paragraph()   #添加题干
    para.add_run(str(lst[index][0])+'.\t')  #添加题目的序号
    para.add_run(lst[index][3])

    #添加选项:
    answer_A=doc.add_paragraph()
    answer_A.add_run('A.\t'+str(lst[index][6]))   #答案A

    answer_B = doc.add_paragraph()
    answer_B.add_run('B.\t' + str(lst[index][7]))  # 答案B

    if lst[index][8]!=None:
        answer_C = doc.add_paragraph()
        answer_C.add_run('C.\t' + str(lst[index][8]))  # 答案C

    if lst[index][9] != None:
        answer_D = doc.add_paragraph()
        answer_D.add_run('D.\t' + str(lst[index][9]))  # 答案D


    #如果是多选题的话，可能会有E ,F ,G 选项
    if lst[index][10]!=None:  #说明多选题
        answer_E = doc.add_paragraph()
        answer_E.add_run('E.\t' + str(lst[index][10]))  # 答案E

    if lst[index][11]!=None:
        answer_F = doc.add_paragraph()
        answer_F.add_run('F.\t' + str(lst[index][11]))  # 答案F

    if lst[index][12] != None:
        answer_G = doc.add_paragraph()
        answer_G.add_run('G.\t' + str(lst[index][12]))  # 答案G

    #正确答案

    right_answer=doc.add_paragraph()
    right_answer.add_run('正确答案为:\t'+str(lst[index][5])).font.color.rgb=RGBColor(255,0,0)

    #题与题之间添加一个空行
    doc.add_paragraph()

doc.save('题库.docx')




