# coding:utf-8
# 教育机构：马士兵教育
# 讲    师：杨淑娟
# 直播时间 8:05
# （4）Word中的表格插入图片
#   用到docx  -->Python操作word 安装方式 pip install  python-docx
# 导入
# 循环将图片插入表格中
import os
from docx import Document
from docx.shared import  Cm
lstname=os.listdir() # 获取当前路径下所有文件名(文件夹的名称)
picnames=[]
# 打开Word文件 （创建doc文档对象）
doc=Document('插入图片.docx')
# # 文档中的表格按照序号排列的，第一个表格为0，第二个表格为1....
tb=doc.tables[0]
for i in lstname:
    if i.endswith('.png')or i.endswith('.jpg'):
        picnames.append(i)
# row=1
# col=3
for index,item in enumerate(picnames): # 获取索引与内容
    run = tb.cell(index+1, 3).paragraphs[0].add_run()  # 单元格添加文本块
    # # 文本块中插入图片
    pic=run.add_picture(item)
    #
    # # 设置图片的大小
    pic.height=Cm(4)
    pic.width=Cm(6)
    # row+=1  # 行数要自加，列数不变


doc.save('插入图片.docx')

# # 获取表格的行数与列数，
# print('表格的行数',len(tb.rows))
# print('表格的列数',len(tb.columns))
# # 表格的行与表格的列都是有序号的，第一行为0，第二行为1....
# #   第一列为0，第二列为1，。。。。。
# # 想在表格中，原图片的位置插入图片，需要计算，插入图片的位置是第几行第几列
# # 先获取第0行，行0列的值  （第一行，第一列）
# print(tb.cell(0,0).paragraphs[0].text)
# # 获取待操作的单元素  （1，3） 第二行，第4列
# run=tb.cell(1,3).paragraphs[0].add_run() # 单元格添加文本块
# # 文本块中插入图片
# pic=run.add_picture('picture.png')
#
# # 设置图片的大小
# pic.height=Cm(4)
# pic.width=Cm(6)
# # 保存
